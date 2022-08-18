#By Corey Edh

import re
import docx
#from docx.enum.style import WD_STYLE_TYPE


class Word_Operations():
    def __init__(self, docfileName):
        docfile = docx.Document(docfileName)
        self.save = None  #OutPut File
        self.docfile = docfile
        self.path = None
        self.Sections = None

        

    #Pages to Heading Level 6
    def pagesToHeading(self, start):
        page = int(start)
        paragraphs = self.docfile.paragraphs
        prevpage = 0
        print(page)

        for paragraph in paragraphs:
            if paragraph.text == str(prevpage):
                print("duplicate found:", prevpage)

            if paragraph.text == str(page):
                paragraph.style = self.docfile.styles["Heading 6"]
                prevpage+=1
                print(page)
                page+=1

        print("Pages have been converted")

        #Saves the changes to the output
        self.docfile.save(self.save)

        #Output file becomes main file (Prevents user from overwriting it)
        self.docfile = docx.Document(self.save)
    


        
    def getSections(self, styleOfSectionText):
        section = {}
        doc = self.docfile
        paragraphs = self.docfile.paragraphs

        for text in styleOfSectionText:
            style = text[1]
            section[text[0]] = style

        for paragraph in paragraphs:
            paragraphText = re.sub(r"[\n\t\s]*", "", paragraph.text.lower()) #Remove whitespaces
            if paragraphText in section:
                print(paragraph.text)
                if section[paragraphText] == "H1":
                    paragraph.style = doc.styles["Heading 1"]

                elif section[paragraphText] == "H2":
                    paragraph.style = doc.styles["Heading 2"]

                elif section[paragraphText] == "H3":
                    paragraph.style = doc.styles["Heading 3"]

                elif section[paragraphText] == "H4":
                    paragraph.style = doc.styles["Heading 4"]

                elif section[paragraphText] == "H5":
                    paragraph.style = doc.styles["Heading 5"]

                elif section[paragraphText] == "H6":
                    paragraph.style = doc.styles["Heading 6"]
                
        self.docfile.save(self.save)
        print("Sections has been converted")



    #if image is found, adds new paragraph "Alt text needed"
    def findImagesThatNeedAltText(self):
        paragraphs = self.docfile.paragraphs
        x = 1
    
        for paragraph in paragraphs:
            if 'graphicData' in paragraph._p.xml:
                print(x)
                newParagraph = self.docfile.add_paragraph("Alt text needed", style = "Heading 7")
                paragraph._p.addnext(newParagraph._p)
                x+=1

        self.docfile.save(self.save)
        self.docfile = docx.Document(self.save)
        print("Alt Text reminder have been added next to images")

